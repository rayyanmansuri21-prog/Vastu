from django.shortcuts import render, redirect
from django.shortcuts import get_object_or_404
from .models import Project, UserProfile, PaymentOrder
from .forms import ProjectForm
from django.http import JsonResponse
import base64
from django.http import HttpResponse, JsonResponse
from fpdf import FPDF
from io import BytesIO
from PIL import Image
from django.views.decorators.csrf import csrf_exempt
import random
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.utils.html import format_html
from django.http import HttpResponseRedirect
from django.urls import reverse
from django.contrib.auth import login, get_user_model, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib import messages
import math
import json
from .utils import calculate_zonal_areas
import matplotlib
from .utils import generate_graph_dxf
matplotlib.use('Agg')
from matplotlib import pyplot as plt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader
import io, base64
from django.db.models import Count, OuterRef, Subquery
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from django.core.files.base import ContentFile
import traceback
from .models import Project
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import numpy as np
import matplotlib.pyplot as plt
import io
import json
import base64
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import ezdxf
from django.conf import settings
import tempfile
import google.generativeai as genai
from google.generativeai.types import (HarmCategory, HarmBlockThreshold)
import logging
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
from openai import OpenAI
from django.contrib.auth.decorators import login_required
from .models import VastuReport, Project
from django.views.decorators.http import require_POST
import razorpay
import hmac
import hashlib
from django.db.models import Prefetch, Sum


# razorpay function
PLAN_CONFIG = {
    'basic':    {'amount': 4900,  'extra_projects': 1,  'label': '1 Extra Project'},
    'standard': {'amount': 29900, 'extra_projects': 2,  'label': '2 Extra Projects'},
    'premium':  {'amount': 49900, 'extra_projects': 7,  'label': '7 Extra Projects'},
}
@login_required
@require_POST
def create_payment_order(request):
    try:
        body = json.loads(request.body)
        plan = body.get('plan')

        if plan not in PLAN_CONFIG:
            return JsonResponse({'error': 'Invalid plan'}, status=400)

        config = PLAN_CONFIG[plan]

        client = razorpay.Client(
            auth=(settings.RAZORPAY_KEY_ID, settings.RAZORPAY_KEY_SECRET)
        )

        # Razorpay order create karo
        order_data = {
            'amount': config['amount'],
            'currency': 'INR',
            'payment_capture': 1,
            'notes': {
                'plan': plan,
                'user_id': str(request.user.id),
                'username': request.user.username,
            }
        }
        razorpay_order = client.order.create(data=order_data)

        # DB mein save karo
        PaymentOrder.objects.create(
            user=request.user,
            plan=plan,
            razorpay_order_id=razorpay_order['id'],
            amount=config['amount'],
            extra_projects_granted=config['extra_projects'],
        )

        return JsonResponse({
            'order_id': razorpay_order['id'],
            'amount': config['amount'],
            'currency': 'INR',
            'key': settings.RAZORPAY_KEY_ID,
            'plan_label': config['label'],
            'user_name': request.user.get_full_name() or request.user.username,
            'user_email': request.user.email,
        })

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)
@csrf_exempt
@require_POST
def verify_payment(request):
    try:
        body = json.loads(request.body)
        razorpay_order_id   = body.get('razorpay_order_id')
        razorpay_payment_id = body.get('razorpay_payment_id')
        razorpay_signature  = body.get('razorpay_signature')

        # Signature verify karo
        key_secret = settings.RAZORPAY_KEY_SECRET.encode()
        msg = f"{razorpay_order_id}|{razorpay_payment_id}".encode()
        generated_signature = hmac.new(key_secret, msg, hashlib.sha256).hexdigest()

        if generated_signature != razorpay_signature:
            return JsonResponse({'error': 'Invalid signature'}, status=400)

        # Order fetch karo
        order = PaymentOrder.objects.get(razorpay_order_id=razorpay_order_id)

        if order.status == 'paid':
            return JsonResponse({'message': 'Already processed'})

        # Payment mark karo
        order.razorpay_payment_id = razorpay_payment_id
        order.status = 'paid'
        order.save()

        # User ka project limit badhao
        profile, _ = UserProfile.objects.get_or_create(user=order.user)
        profile.extra_projects += order.extra_projects_granted
        profile.save()

        return JsonResponse({
            'success': True,
            'new_limit': profile.total_limit(),
            'extra_granted': order.extra_projects_granted,
        })

    except PaymentOrder.DoesNotExist:
        return JsonResponse({'error': 'Order not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@login_required
def get_project_status(request):
    profile, _ = UserProfile.objects.get_or_create(user=request.user)
    from .models import Project  # apna Project model import karo
    used = Project.objects.filter(user=request.user).count()
    return JsonResponse({
        'limit': profile.total_limit(),
        'used': used,
        'remaining': profile.total_limit() - used,
    })
#razorpay functions end here

@require_POST
@csrf_exempt
def vastu_ai_suggest(request):
    try:
        body = json.loads(request.body)
        room_list = body.get('roomList', 'No rooms placed yet')
        has_walls = body.get('hasWalls', False)

        client = Groq(api_key=settings.GROQ_API_KEY)

        prompt = f"""You are a Vastu Shastra expert helping a user design their home layout.

Current layout state:
- Walls drawn: {'Yes' if has_walls else 'No walls drawn yet'}
- Rooms placed: {room_list}

Based on traditional Vastu Shastra principles, give 3-4 short, specific, actionable suggestions for room placement.
Each suggestion should follow this format: "→ [Room name]: [Direction/corner] — [1 short reason]"
Be specific about directions (Northeast, Southwest, Southeast, etc.).
Keep total response under 120 words.
Do NOT use markdown, asterisks, or bullet symbols. Use → as the only bullet."""

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
        )

        suggestion = response.choices[0].message.content
        return JsonResponse({'suggestion': suggestion})

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

# function for AI generated Report
@login_required
def generate_vastu_report(request, project_id):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    try:
        project = Project.objects.get(id=project_id, user=request.user)
        data = json.loads(request.body)

        # Frontend se aaya data
        degree_north = data.get('degree_north', 0)
        degree_east = data.get('degree_east', 0)
        degree_south = data.get('degree_south', 0)
        degree_west = data.get('degree_west', 0)
        compass_degree = data.get('compass_degree', 0)
        zone_percentages = data.get('zone_percentages', {})  # optional

        # Claude ke liye structured prompt
        prompt = f"""
You are an expert Vastu Shastra consultant with 20+ years of experience.
Analyze the following blueprint data and provide a detailed Vastu report.

## Blueprint Data:
- Property compass orientation (North degree): {degree_north}°
- East degree: {degree_east}°
- South degree: {degree_south}°
- West degree: {degree_west}°
- Compass rotation applied: {compass_degree}°

## Zone Area Percentages (if available):
{json.dumps(zone_percentages, indent=2) if zone_percentages else "Not provided - analyze based on compass directions only."}

## Your Task:
Provide a comprehensive Vastu analysis. Respond ONLY in this exact JSON format, no extra text:

{{
  "vastu_score": <integer 0-100>,
  "overall_summary": "<2-3 sentence overall assessment>",
  "zone_analysis": {{
    "North":  {{ "element": "Water", "ruling_deity": "Kuber", "ideal_use": "...", "current_status": "Good/Average/Poor", "impact": "..." }},
    "NorthEast": {{ "element": "Space", "ruling_deity": "Ishaan", "ideal_use": "...", "current_status": "Good/Average/Poor", "impact": "..." }},
    "East":   {{ "element": "Air", "ruling_deity": "Indra", "ideal_use": "...", "current_status": "Good/Average/Poor", "impact": "..." }},
    "SouthEast": {{ "element": "Fire", "ruling_deity": "Agni", "ideal_use": "...", "current_status": "Good/Average/Poor", "impact": "..." }},
    "South":  {{ "element": "Earth", "ruling_deity": "Yama", "ideal_use": "...", "current_status": "Good/Average/Poor", "impact": "..." }},
    "SouthWest": {{ "element": "Earth", "ruling_deity": "Nairutya", "ideal_use": "...", "current_status": "Good/Average/Poor", "impact": "..." }},
    "West":   {{ "element": "Water", "ruling_deity": "Varuna", "ideal_use": "...", "current_status": "Good/Average/Poor", "impact": "..." }},
    "NorthWest": {{ "element": "Air", "ruling_deity": "Vayu", "ideal_use": "...", "current_status": "Good/Average/Poor", "impact": "..." }}
  }},
  "positive_aspects": ["<aspect 1>", "<aspect 2>", "<aspect 3>"],
  "doshas": [
    {{ "name": "<dosha name>", "direction": "<direction>", "severity": "Minor/Major", "description": "<what is wrong>" }}
  ],
  "recommendations": [
    {{ "priority": "High/Medium/Low", "direction": "<direction>", "action": "<what to do>", "benefit": "<expected result>" }}
  ]
}}
"""

        # Claude API call
        client = OpenAI(
            api_key=settings.GROQ_API_KEY,
            base_url="https://api.groq.com/openai/v1"
        )

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system",
                 "content": "You are an expert Vastu Shastra consultant. Always return valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )

        raw_text = response.choices[0].message.content.strip()

        # JSON parse karo safely
        try:
            report_data = json.loads(raw_text)
        except json.JSONDecodeError:
            # Agar JSON wrapped ho to clean karo
            import re
            match = re.search(r'\{.*\}', raw_text, re.DOTALL)
            if match:
                report_data = json.loads(match.group())
            else:
                return JsonResponse({'error': 'AI response could not be parsed.'}, status=500)

        # DB mein save karo (update or create)
        report, created = VastuReport.objects.update_or_create(
            project=project,
            defaults={
                'vastu_score': report_data.get('vastu_score', 0),
                'overall_summary': report_data.get('overall_summary', ''),
                'zone_analysis': report_data.get('zone_analysis', {}),
                'positive_aspects': report_data.get('positive_aspects', []),
                'doshas': report_data.get('doshas', []),
                'recommendations': report_data.get('recommendations', []),
            }
        )

        return JsonResponse({'success': True, 'report': {
            'vastu_score': report.vastu_score,
            'overall_summary': report.overall_summary,
            'zone_analysis': report.zone_analysis,
            'positive_aspects': report.positive_aspects,
            'doshas': report.doshas,
            'recommendations': report.recommendations,
        }})

    except Project.DoesNotExist:
        return JsonResponse({'error': 'Project not found'}, status=404)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)



def calculate_polygon_centroid(points):
    area = 0
    Cx = 0
    Cy = 0
    n = len(points)

    for i in range(n):
        x1, y1 = points[i]
        x2, y2 = points[(i + 1) % n]

        cross = x1 * y2 - x2 * y1
        area += cross

        Cx += (x1 + x2) * cross
        Cy += (y1 + y2) * cross

    area *= 0.5

    if area == 0:
        return None

    Cx /= (6 * area)
    Cy /= (6 * area)

    return Cx, Cy
@csrf_exempt
def calculate_center(request):
    if request.method == "POST":
        data = json.loads(request.body)
        points = data.get("points", [])

        if len(points) < 3:
            return JsonResponse({"error": "Invalid polygon"})

        centroid = calculate_polygon_centroid(points)

        if centroid is None:
            return JsonResponse({"error": "Area zero"})

        return JsonResponse({
            "center_x": centroid[0],
            "center_y": centroid[1]
        })


if "centroid" not in locals():
    centroid = {"x": 0, "y": 0}


def generate_otp():
    return str(random.randint(1000, 9999))  # 4-digit OTP


# FUNCTIONS FOT CHATBOT-----------------------------------------------
import os
import json
from dotenv import load_dotenv
from groq import Groq
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

load_dotenv()

@csrf_exempt
def chat_api(request):
    if request.method != "POST":
        return JsonResponse({"error": "Invalid request"}, status=405)

    try:
        data = json.loads(request.body)
        user_message = data.get("message", "").strip()
    except:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    if not user_message:
        return JsonResponse({"error": "Empty message"}, status=400)

    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return JsonResponse({"error": "API key missing"}, status=500)

    try:
        client = Groq(api_key=api_key)

        completion = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional Vastu Shastra consultant. Give clear practical advice."
                },
                {
                    "role": "user",
                    "content": user_message
                }
            ]
        )

        reply = completion.choices[0].message.content

        return JsonResponse({"reply": reply})

    except Exception as e:
        return JsonResponse({"error": str(e)}, status=500)
# other functionss starts from here----------------------------------

def knowledge(request):
    return render(request, 'knowledge.html')

def application_insights(request):
    return render(request, 'application_insights.html')

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        email = request.POST.get('email')

        if username and email:
            otp = generate_otp()

            # ✨ HTML message for email
            html_message = f"""
            <div style="font-family: Arial, sans-serif; max-width: 500px; margin: auto; border: 1px solid #ddd; padding: 20px; border-radius: 10px; background-color: #f9f9f9;">
                <h2 style="color: #4A00E0;">🔮 Cosmo-Vastu OTP Verification</h2>
                <p>Hi <strong>{username}</strong>,</p>
                <p>Thank you for signing up at <strong>Cosmo Vastu</strong>!<br>
                Use the following OTP to complete your verification:</p>
                <div style="font-size: 24px; font-weight: bold; color: #4A00E0; margin: 20px 0; letter-spacing: 5px;">
                    {otp}
                </div>
                <p>This OTP is valid for a limited time only. Please do not share it with anyone.</p>
                <br>
                <p style="font-size: 13px; color: #777;">© {username} | Cosmo Vastu Team</p>
            </div>
            """

            send_mail(
                subject='Your Cosmo-Vastu OTP',
                message=f'Hello {username}, your OTP is: {otp}',  # fallback plain text
                from_email='your_email@gmail.com',
                recipient_list=[email],
                fail_silently=False,
                html_message=html_message
            )

            request.session['otp'] = otp
            request.session['email'] = email
            request.session['username'] = username

            return redirect('verify_otp')
        else:
            return render(request, 'login.html', {'error': 'All fields are required.'})

    return render(request, 'login.html')


def verify_otp(request):
    # 1️⃣  Handle POST – जब यूज़र ने 4 डिजिट लिखकर सबमिट किया
    if request.method == "POST":
        # 4 textbox values → एक स्ट्रिंग
        otp_digits = [request.POST.get(f'otp{i}') for i in range(1, 5)]
        if None in otp_digits:  # कोई बॉक्स खाली है?
            return render(request, "verify_otp.html",
                          {"error": "OTP incomplete."})

        user_otp = "".join(otp_digits)  # e.g. "5739"

        # 2️⃣  Mail भेजते वक़्त जो चीज़ें सेशन में रखी थीं, निकाल लो
        session_otp = request.session.get("otp")
        email = request.session.get("email")
        username = request.session.get("username")

        # 3️⃣  OTP मैच करता है?
        if user_otp == session_otp:
            #  ⬇️ auth_user में यूज़र ढूँढो या नया बना दो
            user, created = User.objects.get_or_create(
                username=username,
                defaults={"email": email},
            )
            if created:
                user.set_unusable_password()  # हम OTP लॉग‑इन कर रहे हैं
                user.save()

            # 4️⃣  लॉग‑इन कर दो (session auth)
            login(request, user)

            # 5️⃣  Email को फिर से session में डाल दो (आगे काम आयेगा)
            request.session["email"] = user.email

            return redirect("dashboard")

        # ❌ OTP ग़लत
        return render(request, "verify_otp.html",
                      {"error": "Invalid OTP. Try again."})

    # GET request – सिर्फ़ पेज दिखा दो
    return render(request, "verify_otp.html")


@login_required
def dashboard_view(request):
    user = request.user
    recent_projects = Project.objects.filter(user=user).order_by('-created_at')[:3]

    # Fetch project limit from user profile
    try:
        user_profile = UserProfile.objects.get(user=user)
        project_limit = user_profile.project_limit
    except UserProfile.DoesNotExist:
        project_limit = 5  # fallback in case profile not found

    current_project_count = Project.objects.filter(user=user).count()
    limit_reached = current_project_count >= project_limit

    all_projects = Project.objects.filter(user=user).order_by('-created_at')

    return render(request, 'dashboard.html', {
        'recent_projects': recent_projects,
        'limit_reached': limit_reached,
        'user_profile': user,
        'all_projects': all_projects,
        'project_count': current_project_count,
        'project_limit': project_limit
    })


def success_view(request):
    return render(request, 'success.html')


def project_success(request):
    return render(request, 'blueprint_workspace.html')


# @login_required
# def create_project(request):
#     if request.method == 'POST':
#         user = request.user
#         user_profile = request.user.userprofile
#         user_projects = Project.objects.filter(user=request.user).count()

#         if user_projects >= user_profile.project_limit:
#             messages.error(request, "You have reached your project creation limit.")
#             return redirect('dashboard')

#         name = request.POST.get('projectName')
#         description = request.POST.get('description')
#         status = request.POST.get('status')
#         category = request.POST.get('category')
#         blueprint = request.FILES.get('blueprint')

#         if not all([name, description, status, category, blueprint]):
#             messages.error(request, "All fields are required.")
#             return redirect('dashboard')

#         Project.objects.create(
#             user=user,
#             name=name,
#             description=description,
#             status=status,
#             category=category,
#             blueprint=blueprint,
#         )

#         return redirect('blueprint_workspace', project_id=project.id)

#     return redirect('dashboard')

# old function 3 project creation limit
@login_required
def create_project(request):
    if request.method == 'POST':
        user = request.user
        user_profile = UserProfile.objects.get(user=request.user)
        project_count = Project.objects.filter(user=request.user).count()
        if project_count >= user_profile.project_limit:
            messages.error(request, "You have reached your project creation limit Please Upgrade to Premium to continue.")
            return redirect('dashboard')

        # --- grab the fields that came from the form ---
        name = request.POST.get('projectName')
        description = request.POST.get('description')
        status = request.POST.get('status')
        category = request.POST.get('category')
        blueprint = request.FILES.get('blueprint')

        if not all([name, description, status, category, blueprint]):
            messages.error(request, "All fields are required.")
            return redirect('dashboard')

        project = Project.objects.create(
            user=user,  # 🔗 who owns it
            name=name,
            description=description,
            status=status,
            category=category,
            blueprint=blueprint,
        )
        return redirect('blueprint_workspace', project_id=project.id)

    # GET
    return redirect('dashboard')

@login_required
def create_project_from_layout(request):
    if request.method != 'POST':
        return redirect('dashboard')

    user = request.user

    # Project limit check — same as existing create_project
    try:
        user_profile = UserProfile.objects.get(user=user)
    except UserProfile.DoesNotExist:
        user_profile = None

    if user_profile:
        project_count = Project.objects.filter(user=user).count()
        if project_count >= user_profile.project_limit:
            messages.error(request, "You have reached your project creation limit.")
            return redirect('dashboard')

    name        = request.POST.get('projectName', '').strip()
    description = request.POST.get('description', '').strip()
    status      = request.POST.get('status', '').strip()
    category    = request.POST.get('category', '').strip()
    image_data  = request.POST.get('layout_image_data', '')

    if not all([name, description, status, category, image_data]):
        messages.error(request, "All fields are required.")
        return redirect('create_layout')

    try:
        # Base64 image → file
        if ',' in image_data:
            image_data = image_data.split(',')[1]

        image_bytes  = base64.b64decode(image_data)
        image_file   = ContentFile(image_bytes, name=f'layout_{user.id}_{name[:20]}.png')

        project = Project.objects.create(
            user=user,
            name=name,
            description=description,
            status=status,
            category=category,
            blueprint=image_file,
        )

        return redirect('blueprint_workspace', project_id=project.id)

    except Exception as e:
        messages.error(request, f"Error creating project: {str(e)}")
        return redirect('create_layout')


@login_required
def create_layout(request):
    return render(request, 'create_layout.html')


def blueprint_workspace(request, project_id):
    project = get_object_or_404(Project, id=project_id, user=request.user)
    return render(request, 'blueprint_workspace.html', {'project': project, 'project_id': project_id})


# def blueprint_workspace(request, blueprint_path):
#     return render(request, 'blueprint_workspace.html', {'blueprint_path': blueprint_path})

# @login_required
def increase_project_limit(request, user_id):
    if request.method == "POST":
        extra_limit = int(request.POST.get("extra_limit", 0))

        user_to_update = get_object_or_404(User, id=user_id)

        user_profile, created = UserProfile.objects.get_or_create(user=user_to_update)

        # Limit ko badhayein
        user_profile.project_limit += extra_limit
        user_profile.save()

        # User ko email bhejein
        email_context = {
            'username': user_to_update.username,
            'extra_limit': extra_limit,
            'new_limit': user_profile.project_limit,
        }

        # Step 3: HTML template ko render karein
        # 'emails/project_limit_increased.html' path se template ko load karein aur context ke saath render karein.
        html_message = render_to_string('emails/project_limit_increased.html', email_context)

        # Step 4: HTML message ka plain text version banayein
        # Yeh un email clients ke liye fallback hai jo HTML support nahi karte.
        plain_message = strip_tags(html_message)

        # Step 5: User ko naye design wala email bhejein
        try:
            send_mail(
                subject="🚀 Project Limit Increased",
                message=plain_message,  # Plain text message (fallback ke liye)
                from_email="your_admin_email@example.com",  # Apna email yahan daalein
                recipient_list=[user_to_update.email],
                fail_silently=False,
                html_message=html_message,  # HTML version yahan pass karein
            )
            messages.success(request, f"{user_to_update.username}'s project limit has been increased successfully.")
        except Exception as e:
            messages.warning(request,
                             f"Limit was increased for {user_to_update.username}, but failed to send email. Error: {e}")

        # Apne admin dashboard ke URL name par redirect karein
        return redirect("admin_dashboard")

        # Agar POST request nahi hai, to bhi dashboard par bhej dein
    return redirect("admin_dashboard")


# @csrf_exempt
# def download_blueprint(request):
#     if request.method == "POST":
#         # 1. Image
#         image_data = request.POST.get("image_data", "").split(",")[-1]
#         blueprint_img = base64.b64decode(image_data)

#         # 2. Session values
#         grid_points = request.session.get("grid_points", [])
#         compass_center = request.session.get("compass_center", [0, 0])
#         divisions = request.session.get("divisions", 8)

#         # 3. Calculate direction-wise area
#         direction_counts = calculate_directional_areas(
#             [tuple(p) for p in grid_points], tuple(compass_center), divisions
#         )
#         zones = divisions
#         values = list(direction_counts.values())

#         if zones > 0 and values:
#             total = sum(values)
#             avg = total / zones
#             max_line = (avg + max(values)) / 2
#             min_line = (avg + min(values)) / 2
#         else:
#             total = avg = max_line = min_line = 0

#         # 4. Assign zone-based color scheme
#         color_map = {}
#         direction_keys =  list(direction_counts.keys())
#         if divisions == 8:
#             color_map = {
#                 "N": "blue", "NE": "blue", "E": "green", "SE": "red",
#                 "S": "red", "SW": "yellow", "W": "grey", "NW": "grey"
#             }
#         elif divisions == 16:
#             for d in direction_keys:
#                 if d in ["NNW", "N", "NNE", "NE"]:
#                     color_map[d] = "blue"
#                 elif d in ["ENE", "E", "ESE"]:
#                     color_map[d] = "green"
#                 elif d in ["SE", "SSE", "S"]:
#                     color_map[d] = "red"
#                 elif d in ["SW", "SSW"]:
#                     color_map[d] = "yellow"
#                 elif d in ["WSW", "W", "WNW", "NW"]:
#                     color_map[d] = "grey"
#                 else:
#                     color_map[d] = "black"
#         elif divisions == 32:
#             for d in direction_keys:
#                 if d in ["N5", "N6", "N7", "N8", "E1", "N2", "N3", "N4"]:
#                     color_map[d] = "blue"
#                 elif d in ["E2", "E3", "E4", "E5", "E6", "E7"]:
#                     color_map[d] = "green"
#                 elif d in ["E8", "S1", "S2", "S3", "S4", "S5"]:
#                     color_map[d] = "red"
#                 elif d in ["S6", "S7", "S8", "W1"]:
#                     color_map[d] = "yellow"
#                 elif d in ["W2", "W3", "W4", "W5", "W6", "W7", "W8", "N1"]:
#                     color_map[d] = "grey"
#                 else:
#                     color_map[d] = "black"


#         bar_colors = [color_map.get(k, "black") for k in direction_counts.keys()]

#         # 5. Plot Graph
#         fig, ax = plt.subplots(figsize=(10, 4))
#         ax.bar(direction_counts.keys(), direction_counts.values(), color=bar_colors)
#         ax.axhline(avg, color='red', linestyle='--', label='AVG AREA')
#         ax.axhline(max_line, color='purple', linestyle='--', label='MAX LINE')
#         ax.axhline(min_line, color='green', linestyle='--', label='MIN LINE')
#         ax.set_title("Zone-wise Area Distribution with Threshold Indicators")
#         ax.set_xlabel("Zone")
#         ax.set_ylabel("Area (sq units)")
#         ax.legend()
#         ax.grid(True)
#         plt.xticks(rotation=45)
#         plt.tight_layout()

#         graph_buffer = io.BytesIO()
#         plt.savefig(graph_buffer, format='PNG')
#         graph_buffer.seek(0)

#         # 6. Create PDF
#         pdf_buffer = io.BytesIO()
#         p = canvas.Canvas(pdf_buffer, pagesize=A4)

#         # First page - Blueprint
#         p.drawString(50, 800, "Cosmo Vastu - Blueprint Analysis")
#         blueprint_reader = ImageReader(io.BytesIO(blueprint_img))
#         p.drawImage(blueprint_reader, 50, 200, width=500, preserveAspectRatio=True, mask='auto')

#         p.setFont("Helvetica", 10)
#         p.drawString(50, 180, f"Total Area: {total}")
#         p.drawString(200, 180, f"Zones: {zones}")
#         p.drawString(50, 160, f"Avg Area: {round(avg, 2)}")
#         p.drawString(200, 160, f"Max Line: {round(max_line, 2)}")
#         p.drawString(350, 160, f"Min Line: {round(min_line, 2)}")

#         p.showPage()

#         # Second page - Graph
#         p.drawImage(ImageReader(graph_buffer), 50, 250, width=500, preserveAspectRatio=True, mask='auto')
#         p.showPage()

#         p.save()
#         pdf_buffer.seek(0)
#         return HttpResponse(pdf_buffer, content_type="application/pdf")

# -----------------------------------------------------------------------------
# DXF Conversion Helper Function (Corrected)
# -----------------------------------------------------------------------------
import ezdxf
import os
import tempfile
import io
import base64
import json
import math
import traceback
from django.core.files.base import ContentFile
from django.shortcuts import get_object_or_404
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt


# >>> आपको अपनी Django model को यहां आयात करना होगा:
# from .models import Project 

# ==============================================================================
# I. DXF FILE FETCHING (Database interaction logic)
# ==============================================================================

def fetch_dxf_for_division(project_id, divisions):
    """
    Fetches the appropriate DXF file for the given project.
    """
    project = get_object_or_404(Project, pk=project_id)

    if divisions == 8:
        dxf_file_field = project.divided_8_dxf
    elif divisions == 16:
        dxf_file_field = project.divided_16_dxf
    elif divisions == 32:
        dxf_file_field = project.divided_32_dxf
    else:
        raise ValueError(f"Invalid divisions value: {divisions}")

    if dxf_file_field and dxf_file_field.storage.exists(dxf_file_field.name):
        dxf_file_field.open('rb')
        dxf_bytes = dxf_file_field.read()
        dxf_file_field.close()

        # ✅ DEBUG: Check what type of data we're getting
        print(f"DEBUG: DXF data type: {type(dxf_bytes)}")
        print(f"DEBUG: DXF data length: {len(dxf_bytes) if dxf_bytes else 0}")

        return dxf_bytes
    else:
        raise FileNotFoundError(f"DXF file for divisions {divisions} not found in database.")


# ==============================================================================
# II. DXF PARSING & DIRECTIONAL CALCULATION (AutoCAD Logic)
# ==============================================================================
# new code from deepseek*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-*-
# ---- Add these helper wrappers (place after process_dxf_manually) ----
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt

from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from .utils import calculate_zonal_areas
import math, ezdxf, io, base64, json, traceback


# main/views.py

# ... (previous code) ...

@csrf_exempt
def graph_preview(request):
    """Return PNG preview of the graph in a new tab, using session zone_measurements."""
    try:
        divisions = int(request.GET.get("divisions", request.session.get("divisions", 8)))
        zone_measurements = request.session.get("zone_measurements", [])
        if not zone_measurements:
            return HttpResponse("No zone measurements in session. Run analysis first.", status=400)
        area_results = calculate_zonal_areas(zone_measurements)
        direction_areas = {name: info.get('area_sq_inches', info.get('area_sq_ft', 0))
                           for name, info in area_results['zonal_areas'].items()}
        from .utils import harmonize_direction_pairs
        direction_areas = harmonize_direction_pairs(direction_areas, divisions)
        from .utils import generate_graph_png
        img = generate_graph_png(direction_areas, divisions, unit_label='sq in')
        return HttpResponse(img.getvalue(), content_type='image/png')
    except Exception as e:
        traceback.print_exc()
        return HttpResponse(f"Error: {e}", status=500)


@csrf_exempt
def plot_graph_and_area(request):
    """
    Generates DXF bar graph with direction-wise area distribution.
    Returns JSON with base64 encoded DXF data.
    """
    try:
        if request.method != "POST":
            return JsonResponse({'error': 'Only POST allowed'}, status=405)

        # Parse request data
        try:
            body = request.body.decode('utf-8').strip()
            data = json.loads(body) if body else {}
        except json.JSONDecodeError:
            return JsonResponse({'error': 'Invalid JSON'}, status=400)

        divisions = int(data.get("divisions", 8))
        # Directly use zone_measurements from the POST data
        zone_measurements = data.get("zone_measurements", [])

        # --- IMPORTANT: Ensure compass_center and compass_rotation are also passed directly or retrieved accurately ---
        # The frontend `plotGraph` also sends `compass_rotation` and `centroid`.
        # Your `generate_graph_dxf` doesn't directly use these, but `calculate_zonal_areas` doesn't need them.
        # However, `calculate_directional_areas` in `utils.py` does use them.
        # Make sure `generate_graph_dxf` in `utils.py` is robust and the data it uses is accurate.

        if not zone_measurements:
            # If the frontend truly didn't send any, then this is an error from the frontend.
            return JsonResponse({'error': 'No zone measurements provided in the request body.'}, status=400)

        # Calculate areas for all zones
        from .utils import calculate_zonal_areas
        area_results = calculate_zonal_areas(zone_measurements)

        # Prepare direction-wise area mapping (prefer sq inches if available)
        direction_areas = {}
        for zone_name, zone_data in area_results['zonal_areas'].items():
            value = zone_data.get('area_sq_inches')
            if value is None:
                value = zone_data.get('area_sq_ft', 0)
            direction_areas[zone_name] = value
        from .utils import harmonize_direction_pairs
        direction_areas = harmonize_direction_pairs(direction_areas, divisions)

        # Generate DXF bar graph (unit aware)
        from .utils import generate_graph_dxf
        unit_label = 'sq in' if 'area_sq_inches' in next(iter(area_results['zonal_areas'].values())) else 'sq ft'
        dxf_bytes = generate_graph_dxf(direction_areas, divisions, unit_label=unit_label)

        if not dxf_bytes:
            return JsonResponse({'error': 'Failed to generate DXF graph'}, status=500)

        # Return as downloadable file
        response = HttpResponse(dxf_bytes.getvalue(), content_type='application/dxf')
        response['Content-Disposition'] = f'attachment; filename="area_graph_{divisions}_divisions.dxf"'
        return response

    except Exception as e:
        print(f"❌ Error in plot_graph_and_area: {e}")
        traceback.print_exc()
        return JsonResponse({'error': f"Server Error: {str(e)}"}, status=500)


@csrf_exempt
def analyze_grid(request):
    """
    Receives zone measurements from frontend.
    Calculates area for each zone (width x height in inches, converted to sq ft).
    Returns JSON with area calculations.
    """
    try:
        if request.method != "POST":
            return JsonResponse({'error': 'Only POST allowed'}, status=405)

        data = json.loads(request.body)

        print(f"🔍 DEBUG: Received data keys: {data.keys()}")
        print(f"🔍 DEBUG: Full data: {data}")

        # ✅ Handle multiple data formats
        zone_measurements = data.get("zone_measurements", [])

        # If zone_measurements is empty, try alternate field names
        if not zone_measurements:
            zone_measurements = data.get("zones", [])

        # If not provided, derive from grid_data using bounding box per direction and a cell size in inches.
        if not zone_measurements:
            grid_data = data.get("grid_data", [])
            if grid_data:
                divisions = int(data.get("divisions", 8))
                compass_center = data.get("compass_center", [0, 0])
                compass_rotation = float(data.get("compass_rotation", 0))
                cell_inches = float(data.get("cell_inches", 12.0))  # inches per grid step (frontend can send)

                # Normalize points
                pts = []
                for p in grid_data:
                    if isinstance(p, (list, tuple)) and len(p) >= 2:
                        pts.append((float(p[0]), float(p[1])))
                    elif isinstance(p, dict):
                        pts.append((float(p.get("x", 0)), float(p.get("y", 0))))

                if not pts:
                    zone_measurements = []
                else:
                    # infer grid step from unique sorted coords
                    def infer_step(values):
                        vs = sorted(set(values))
                        diffs = [b - a for a, b in zip(vs, vs[1:]) if b - a > 0]
                        return min(diffs) if diffs else 1.0

                    step_x = infer_step([x for x, _ in pts])
                    step_y = infer_step([y for _, y in pts])
                    step = step_x if step_x > 0 else (step_y if step_y > 0 else 1.0)

                    # Labels order
                    if divisions == 8:
                        labels = ["N", "NE", "E", "SE", "S", "SW", "W", "NW"]
                    elif divisions == 16:
                        labels = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE", "S", "SSW", "SW", "WSW", "W", "WNW",
                                  "NW", "NNW"]
                    elif divisions == 32:
                        labels = [
                            "N5", "N6", "N7", "N8",
                            "E1", "E2", "E3", "E4", "E5", "E6", "E7", "E8",
                            "S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8",
                            "W1", "W2", "W3", "W4", "W5", "W6", "W7", "W8",
                            "N1", "N2", "N3", "N4",
                        ]
                    else:
                        labels = [f"Z{i + 1}" for i in range(divisions)]

                    cx, cy = float(compass_center[0]), float(compass_center[1])
                    sector = 360.0 / divisions

                    def bearing(x, y):
                        dx, dy = x - cx, y - cy
                        ang = (90.0 - math.degrees(math.atan2(dy, dx))) % 360.0
                        return (ang + compass_rotation) % 360.0

                    # Group points per direction
                    buckets = {lab: [] for lab in labels}
                    for x, y in pts:
                        idx = int(((bearing(x, y) + sector / 2.0) % 360.0) // sector) % divisions
                        lab = labels[idx]
                        buckets[lab].append((x, y))

                    # Build zone_measurements using bounding box per direction
                    zone_measurements = []
                    for lab in labels:
                        pts_lab = buckets.get(lab, [])
                        if not pts_lab:
                            continue
                        xs = [p[0] for p in pts_lab]
                        ys = [p[1] for p in pts_lab]
                        width_cells = max(1, int(round((max(xs) - min(xs)) / step)) + 1)
                        height_cells = max(1, int(round((max(ys) - min(ys)) / step)) + 1)
                        zone_measurements.append({
                            "zone_name": lab,
                            "width_inches": width_cells * cell_inches,
                            "height_inches": height_cells * cell_inches,
                        })

                print(f"✅ Derived {len(zone_measurements)} zones from grid_data using cell_inches={cell_inches}")

        centroid = data.get("centroid", {"x": 0, "y": 0})
        compass_rotation = float(data.get("compass_rotation", 0))
        compass_center = data.get("compass_center", [centroid.get("x", 0), centroid.get("y", 0)])
        divisions = int(data.get("divisions", len(zone_measurements) or 8))

        if not zone_measurements:
            return JsonResponse({
                'error': 'No zone measurements provided',
                'hint': 'Please send zone_measurements with width_inches and height_inches for each zone'
            }, status=400)

        # Calculate areas using formula: Area (sq inches) = width x height
        area_results = calculate_zonal_areas(zone_measurements)
        direction_areas = {name: info.get('area_sq_inches', info.get('area_sq_ft', 0))
                           for name, info in area_results['zonal_areas'].items()}

        # Save to session for later use
        request.session["zone_measurements"] = zone_measurements
        request.session["centroid"] = centroid
        request.session["compass_rotation"] = compass_rotation
        request.session["compass_center"] = compass_center
        request.session['direction_areas'] = direction_areas
        request.session["divisions"] = divisions
        request.session.modified = True

        print(f"✅ Processing {len(zone_measurements)} zones")

        return JsonResponse({
            "status": "success",
            "area_results": area_results,
            "message": f"Calculated areas for {area_results['total_zones']} zones",
            "divisions": divisions
        })

    except Exception as e:
        print(f"❌ Error in analyze_grid: {e}")
        traceback.print_exc()
        return JsonResponse({'error': f"Server Error: {str(e)}"}, status=500)


@csrf_exempt  # CSRF सुरक्षा को अक्षम करता है, यदि आप POST अनुरोधों के लिए इसका उपयोग कर रहे हैं। उत्पादन में, इसे सही ढंग से प्रबंधित करें।
def display_graph(request, project_id):
    if request.method == 'POST':
        try:
            divisions = int(request.POST.get('divisions', 8))  # डिफ़ॉल्ट रूप से 8 डिवीजन

            # यहां आपको `direction_areas` डेटा प्राप्त करने की आवश्यकता होगी।
            # यह डेटा आमतौर पर `analyze_grid` फ़ंक्शन से आता है और session में सहेजा जाता है।
            # अभी के लिए, हम एक डमी डेटा का उपयोग करेंगे या session से प्राप्त करने का प्रयास करेंगे।
            # **महत्वपूर्ण:** आपको यह डेटा session में सहेजना होगा जब `analyze_grid` चलता है।
            # उदाहरण के लिए, analyze_grid में:
            # request.session['direction_areas'] = direction_counts # या जो भी गणना की जाती है

            direction_areas = request.session.get('direction_areas', {})
            if not direction_areas:
                # यदि session में कोई डेटा नहीं है, तो कुछ डमी डेटा का उपयोग करें या एरर दिखाएं
                print("WARNING: No 'direction_areas' found in session. Using dummy data.")
                direction_areas = {
                    "N": 100, "NE": 150, "E": 200, "SE": 120,
                    "S": 80, "SW": 90, "W": 180, "NW": 130
                }
                if divisions == 16:
                    direction_areas.update({
                        "NNE": 110, "ENE": 160, "ESE": 210, "SSE": 130,
                        "SSW": 70, "WSW": 100, "WNW": 190, "NNW": 140
                    })
                # आप यहां अधिक विस्तृत डमी डेटा जोड़ सकते हैं
                # या एक त्रुटि पृष्ठ पर रीडायरेक्ट कर सकते हैं

            # generate_graph_dxf फ़ंक्शन को कॉल करें
            dxf_data_stream = generate_graph_dxf(direction_areas, divisions)

            if dxf_data_stream:
                # DXF डेटा को base64 में एन्कोड करें ताकि इसे HTML में एम्बेड किया जा सके
                encoded_dxf = base64.b64encode(dxf_data_stream.getvalue()).decode('utf-8')

                context = {
                    'encoded_dxf_data': encoded_dxf,
                    'project_id': project_id,  # यदि आपको इसकी आवश्यकता है
                    'divisions': divisions,  # यदि आपको इसकी आवश्यकता है
                }
                return render(request, 'your_app/graph_display.html', context)
            else:
                return HttpResponse("Error generating DXF graph.", status=500)

        except Exception as e:
            import traceback
            traceback.print_exc()
            return HttpResponse(f"An error occurred: {e}", status=500)
    return redirect('some_default_view')  # यदि GET अनुरोध है तो रीडायरेक्ट करें


def convert_to_dxf(image_file_contentfile, project_id, image_type):
    """
    Create a PROPER DXF with actual entities (grid lines) instead of just an image.
    """
    try:
        import ezdxf
        from io import BytesIO
        from django.core.files.base import ContentFile

        doc = ezdxf.new(dxfversion='R2010')
        msp = doc.modelspace()

        grid_size = 50
        rows = 10
        cols = 10

        for row in range(rows):
            for col in range(cols):
                x = col * grid_size
                y = row * grid_size
                points = [(x, y), (x + grid_size, y),
                          (x + grid_size, y + grid_size), (x, y + grid_size)]
                msp.add_lwpolyline(points, close=True)

        # Add North marker
        center_x = (cols * grid_size) / 2
        center_y = (rows * grid_size) / 2
        msp.add_line((center_x, center_y), (center_x, center_y + 100))
        txt = msp.add_text("N", dxfattribs={'height': 10})
        txt.set_dxf_attrib('insert', (center_x, center_y + 120))

        text_stream = io.StringIO()
        doc.write(text_stream)
        dxf_bytes = text_stream.getvalue().encode('utf-8')

        dxf_filename = f'{project_id}_{image_type}_dxf.dxf'
        dxf_content = ContentFile(dxf_bytes, name=dxf_filename)

        print(f"✅ Created DXF with {rows * cols} entities")
        return dxf_content

    except Exception as e:
        print(f"❌ Error in convert_to_dxf: {e}")
        import traceback
        traceback.print_exc()
        return None


# --- paste this right after your imports in views.py ---

import io
import math
import traceback


# Small helper fallback used by the wrappers
def create_proportional_data(divisions, default_count=60):
    labels = []
    if divisions == 8:
        labels = ["N", "NE", "E", "SE", "S", "SW", "W", "NW"]
    elif divisions == 16:
        labels = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE", "S", "SSW", "SW", "WSW", "W", "WNW", "NW", "NNW"]
    elif divisions == 32:
        labels = [
            "N5", "N6", "N7", "N8",
            "E1", "E2", "E3", "E4", "E5", "E6", "E7", "E8",
            "S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8",
            "W1", "W2", "W3", "W4", "W5", "W6", "W7", "W8",
            "N1", "N2", "N3", "N4"
        ]
    else:
        labels = [f"Z{i + 1}" for i in range(divisions)]
    return {lab: default_count for lab in labels}


def calculate_directional_areas(grid_points_or_dxf, compass_center, divisions, compass_rotation=0):
    """
    Backwards-compatible wrapper: if input is bytes/str, treat as DXF; otherwise treat as list of grid points.
    Returns a dict mapping direction labels to counts.
    """
    try:
        # If bytes/str, parse DXF
        if isinstance(grid_points_or_dxf, (bytes, str)):
            return plot_graph_and_area(request)

        # Otherwise treat as a list of grid points (legacy behavior)
        cx, cy = float(compass_center[0]), float(compass_center[1])
        if divisions == 8:
            labels = ["N", "NE", "E", "SE", "S", "SW", "W", "NW"]
        elif divisions == 16:
            labels = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE", "S", "SSW", "SW", "WSW", "W", "WNW", "NW",
                      "NNW"]
        elif divisions == 32:
            labels = [
                "N5", "N6", "N7", "N8",
                "E1", "E2", "E3", "E4", "E5", "E6", "E7", "E8",
                "S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8",
                "W1", "W2", "W3", "W4", "W5", "W6", "W7", "W8",
                "N1", "N2", "N3", "N4"
            ]
        else:
            labels = [f"Z{i + 1}" for i in range(divisions)]

        direction_counts = {lab: 0 for lab in labels}
        sector_size = 360.0 / float(divisions)

        def angle_from_center(x, y):
            dx = x - cx
            dy = y - cy
            bearing = (90.0 - math.degrees(math.atan2(dy, dx))) % 360.0
            return bearing

        normalized = []
        for p in grid_points_or_dxf:
            if isinstance(p, dict):
                x = float(p.get("x", p.get("0", 0)))
                y = float(p.get("y", p.get("1", 0)))
                inside = p.get("inside", True)
            elif isinstance(p, (list, tuple)):
                if len(p) >= 3:
                    x, y, inside = float(p[0]), float(p[1]), bool(p[2])
                else:
                    x, y, inside = float(p[0]), float(p[1]), True
            else:
                continue
            normalized.append((x, y, inside))

        for (x, y, inside) in normalized:
            if not inside:
                continue
            ang = angle_from_center(x, y)
            ang = (ang + float(compass_rotation)) % 360.0
            idx = int(((ang + sector_size / 2.0) % 360.0) // sector_size) % divisions
            lab = labels[idx] if idx < len(labels) else labels[idx % len(labels)]
            direction_counts[lab] = direction_counts.get(lab, 0) + 1

        return direction_counts

    except Exception as e:
        print("ERROR in calculate_directional_areas fallback:", e)
        traceback.print_exc()
        return create_proportional_data(divisions, default_count=60)


def process_dxf_with_ezdxf(doc, compass_center, divisions, compass_rotation):
    """Process DXF using ezdxf library"""
    msp = doc.modelspace()

    # Label definitions
    labels_8 = ["N", "NE", "E", "SE", "S", "SW", "W", "NW"]
    labels_16 = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE", "S", "SSW", "SW", "WSW", "W", "WNW", "NW", "NNW"]
    labels_32 = [
        "N5", "N6", "N7", "N8", "E1", "E2", "E3", "E4", "E5", "E6", "E7", "E8",
        "S1", "S2", "S3", "S4", "S5", "S6", "S7", "S8", "W1", "W2", "W3", "W4",
        "W5", "W6", "W7", "W8", "N1", "N2", "N3", "N4"
    ]

    if divisions == 8:
        labels = labels_8
    elif divisions == 16:
        labels = labels_16
    elif divisions == 32:
        labels = labels_32
    else:
        labels = [f"S{i + 1}" for i in range(divisions)]

    direction_counts = {label: 0 for label in labels}
    cx, cy = float(compass_center[0]), float(compass_center[1])
    sector_size = 360.0 / divisions

    def get_direction_label(x, y):
        dx, dy = x - cx, y - cy
        bearing = (90.0 - math.degrees(math.atan2(dy, dx))) % 360.0
        bearing = (bearing + compass_rotation) % 360.0
        idx = int(((bearing + sector_size / 2.0) % 360.0) // sector_size) % divisions
        return labels[idx]

    # Process different entity types
    entity_count = 0
    for entity in msp:
        entity_type = entity.dxftype()
        entity_count += 1

        try:
            if entity_type in ["LWPOLYLINE", "POLYLINE"]:
                if entity.is_closed:
                    # Get centroid for closed polylines
                    try:
                        center_x, center_y, _ = entity.boundary_paths.centroid()
                    except:
                        points = list(entity.points())
                        if points:
                            center_x, center_y = points[0][0], points[0][1]
                        else:
                            continue

                    label = get_direction_label(center_x, center_y)
                    direction_counts[label] += 1

            elif entity_type == "INSERT":
                # For block references
                if hasattr(entity.dxf, 'insert'):
                    center_x, center_y = entity.dxf.insert[:2]
                    label = get_direction_label(center_x, center_y)
                    direction_counts[label] += 1

            elif entity_type == "CIRCLE":
                # For circles
                if hasattr(entity.dxf, 'center'):
                    center_x, center_y = entity.dxf.center[:2]
                    label = get_direction_label(center_x, center_y)
                    direction_counts[label] += 1

            elif entity_type == "LINE":
                # For lines, use midpoint
                start_x, start_y = entity.dxf.start[:2]
                end_x, end_y = entity.dxf.end[:2]
                center_x = (start_x + end_x) / 2
                center_y = (start_y + end_y) / 2
                label = get_direction_label(center_x, center_y)
                direction_counts[label] += 1

        except Exception as e:
            print(f"DEBUG: Error processing {entity_type}: {e}")
            continue

    print(f"DEBUG: Processed {entity_count} entities, found {sum(direction_counts.values())} valid ones")

    # If no entities found, use fallback
    if sum(direction_counts.values()) == 0:
        print("⚠️ No entities processed, using fallback")
        return create_fallback_data(divisions)

    return direction_counts


# Function: generate_graph_data_view (Updated to use DXF parsing logic)
# main/views.py

# आवश्यक इम्पोर्ट (सुनिश्चित करें कि आपके पास ये सब हैं)
import json
import base64
import traceback
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt


# --- मान लें कि ये फ़ंक्शन आपके पास मौजूद हैं ---
# from .utils import fetch_dxf_for_division, calculate_directional_areas_from_dxf, generate_graph_dxf
# from .models import Project (या जो भी मॉडल आप उपयोग कर रहे हैं)


@csrf_exempt
def generate_graph_data_view(request, project_id):
    """
    Return a Base64 PNG of the current graph (uses session zone_measurements).
    Frontend expects JSON with key 'graph_dxf_base64' — we keep this name for compatibility.
    """
    if request.method != "POST":
        return JsonResponse({'error': 'Invalid request method! Only POST is allowed.'}, status=405)
    try:
        try:
            data = json.loads(request.body or '{}')
        except json.JSONDecodeError:
            data = {}
        divisions = int(data.get('divisions', request.session.get('divisions', 8)))
        zm = request.session.get('zone_measurements', [])
        if not zm:
            return JsonResponse({'error': 'No zone measurements in session. Run analysis first.'}, status=400)
        area_results = calculate_zonal_areas(zm)
        direction_areas = {k: v.get('area_sq_inches', v.get('area_sq_ft', 0)) for k, v in
                           area_results['zonal_areas'].items()}
        from .utils import harmonize_direction_pairs
        direction_areas = harmonize_direction_pairs(direction_areas, divisions)
        from .utils import generate_graph_png
        png = generate_graph_png(direction_areas, divisions, unit_label='sq in')
        import base64
        b64 = base64.b64encode(png.getvalue()).decode('utf-8')
        return JsonResponse({'graph_dxf_base64': b64})
    except Exception as e:
        traceback.print_exc()
        return JsonResponse({'error': f'Server error: {str(e)}'}, status=500)


# Function: download_blueprint (Updated to use DXF parsing logic)
@csrf_exempt
def download_blueprint(request, project_id):
    """
    Generate and download a DXF bar graph of direction-wise areas.
    Uses zone measurements from request JSON or session.
    """
    if request.method not in ["GET", "POST"]:
        return HttpResponse("This endpoint requires a GET or POST request.", status=405)

    try:
        # Defaults from session
        divisions = int(request.session.get("divisions", 8))
        zone_measurements = request.session.get("zone_measurements", [])

        # Parse JSON body if provided (or regular form POST)
        if request.method == "POST":
            # form-POST support
            if hasattr(request, 'POST'):
                divisions = int(request.POST.get("divisions", divisions))
            # JSON support
            if request.body:
                try:
                    data = json.loads(request.body.decode("utf-8"))
                except json.JSONDecodeError:
                    data = {}
                divisions = int(data.get("divisions", divisions))
                if data.get("zone_measurements"):
                    zone_measurements = data.get("zone_measurements")

        if not zone_measurements:
            return HttpResponse("No zone measurements available. Please analyze grid first.", status=400)

        # Calculate areas and map to directions
        area_results = calculate_zonal_areas(zone_measurements)
        direction_areas = {}
        for name, info in area_results['zonal_areas'].items():
            value = info.get('area_sq_inches')
            if value is None:
                value = info.get('area_sq_ft', 0)
            direction_areas[name] = value
        from .utils import harmonize_direction_pairs
        direction_areas = harmonize_direction_pairs(direction_areas, divisions)

        # Generate DXF graph + keep session updated so preview shows same data
        from .utils import generate_graph_dxf
        unit_label = 'sq in' if 'area_sq_inches' in next(iter(area_results['zonal_areas'].values())) else 'sq ft'
        dxf_bytes_io = generate_graph_dxf(direction_areas, divisions, unit_label=unit_label)
        if not dxf_bytes_io:
            return HttpResponse("Failed to generate DXF graph.", status=500)

        filename = f"directional_graph_{divisions}.dxf"
        response = HttpResponse(dxf_bytes_io.getvalue(), content_type="application/dxf")
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    except Exception as e:
        traceback.print_exc()
        return HttpResponse(f"Server Error while generating DXF: {str(e)}", status=500)


@csrf_exempt  # उत्पादन में, CSRF सुरक्षा को सही ढंग से प्रबंधित करें
def download_graph(request, project_id):
    if request.method == 'POST':
        try:
            divisions = int(request.POST.get('divisions', 8))  # डिफ़ॉल्ट रूप से 8 डिवीजन

            # **महत्वपूर्ण:** यहां आपको `direction_areas` डेटा को session से प्राप्त करने की आवश्यकता होगी।
            # सुनिश्चित करें कि `analyze_grid` फ़ंक्शन इस डेटा को session में सहेजता है।
            # उदाहरण के लिए: request.session['direction_areas'] = your_calculated_areas
            direction_areas = request.session.get('direction_areas', {})

            if not direction_areas:
                print("WARNING: No 'direction_areas' found in session for graph generation. Using dummy data.")
                direction_areas = {
                    "N": 100, "NE": 150, "E": 200, "SE": 120,
                    "S": 80, "SW": 90, "W": 180, "NW": 130
                }
                if divisions == 16:
                    direction_areas.update({
                        "NNE": 110, "ENE": 160, "ESE": 210, "SSE": 130,
                        "SSW": 70, "WSW": 100, "WNW": 190, "NNW": 140
                    })

            from .utils import harmonize_direction_pairs
            direction_areas = harmonize_direction_pairs(direction_areas, divisions)
            dxf_data_stream = generate_graph_dxf(direction_areas, divisions)

            if dxf_data_stream:
                response = HttpResponse(dxf_data_stream.getvalue(), content_type='application/dxf')
                response['Content-Disposition'] = 'attachment; filename="graph_project_{}_divisions_{}.dxf"'.format(
                    project_id, divisions)
                return response
            else:
                return HttpResponse("Failed to generate DXF graph.", status=500)

        except Exception as e:
            import traceback
            traceback.print_exc()
            return HttpResponse(f"An error occurred during DXF generation: {e}", status=500)

    return HttpResponse("Invalid request for graph download.", status=400)


# -----------------------------------------------------------------------------
# DXF Conversion Helper Function (Accepts Image Data and saves it explicitly)
# -----------------------------------------------------------------------------
import ezdxf
import os
import tempfile
import io
import base64
import json
from django.core.files.base import ContentFile
from django.shortcuts import get_object_or_404
from django.http import JsonResponse
from .models import Project  # Assuming your Project model is in models.py


# The save_project_image function remains largely the same, but let's re-include it for completeness
# and ensure the correct variable name is passed to convert_to_dxf
def save_project_image(request, project_id):
    if request.method == 'POST':
        try:
            project = get_object_or_404(Project, pk=project_id)
            data = json.loads(request.body)

            image_data_b64 = data.get('image_data')
            image_type = data.get('image_type')

            if not image_data_b64 or not image_type:
                return JsonResponse({"status": "error", "message": "Missing data"}, status=400)

            format_str, imgstr = image_data_b64.split(';base64,')
            ext = format_str.split('/')[-1]

            # Create ContentFile
            image_file = ContentFile(base64.b64decode(imgstr), name=f'{project_id}_{image_type}.{ext}')

            # Save image to project
            if image_type == 'divided_8':
                project.divided_8_image.save(image_file.name, image_file, save=False)
            elif image_type == 'divided_16':
                project.divided_16_image.save(image_file.name, image_file, save=False)
            elif image_type == 'divided_32':
                project.divided_32_image.save(image_file.name, image_file, save=False)
            elif image_type == 'compass':
                project.compass_image.save(image_file.name, image_file, save=False)
            elif image_type == 'centroid':
                project.centroid_image.save(image_file.name, image_file, save=False)

            project.save()

            # --- DXF Conversion ---
            if image_type in ['divided_8', 'divided_16', 'divided_32']:
                print(f"Creating DXF for {image_type}...")

                # Create a NEW ContentFile for DXF conversion
                image_file_for_dxf = ContentFile(base64.b64decode(imgstr), name=f'{project_id}_{image_type}.{ext}')

                dxf_content = convert_to_dxf(image_file_for_dxf, project_id, image_type)

                if dxf_content:
                    if image_type == 'divided_8':
                        project.divided_8_dxf.save(f"{project_id}_divided_8.dxf", dxf_content, save=False)
                    elif image_type == 'divided_16':
                        project.divided_16_dxf.save(f"{project_id}_divided_16.dxf", dxf_content, save=False)
                    elif image_type == 'divided_32':
                        project.divided_32_dxf.save(f"{project_id}_divided_32.dxf", dxf_content, save=False)

                    project.save()
                    print(f"✅ DXF saved for {image_type}")

            return JsonResponse({"status": "success", "message": f"{image_type} image and DXF saved."})

        except Exception as e:
            print(f"❌ Error in save_project_image: {e}")
            import traceback
            traceback.print_exc()
            return JsonResponse({"status": "error", "message": str(e)}, status=500)

    return JsonResponse({"status": "error", "message": "Invalid request method"}, status=405)


# @csrf_exempt
# def download_blueprint(request):
#     if request.method == "POST":

#         grid_points = request.session.get("grid_points", [])
#         compass_center = request.session.get("compass_center", [0, 0])
#         divisions = request.session.get("divisions", 8)
#         compass_rotation = request.session.get("compass_rotation", 0)

#         # 3. Calculate direction-wise area (Aapka code - Koi badlaav nahi)
#         direction_counts = calculate_directional_areas(
#             [tuple(p) for p in grid_points], tuple(compass_center), divisions, compass_rotation
#         )
#         zones = divisions
#         values = list(direction_counts.values())

#         if zones > 0 and values:
#             total = sum(values)
#             avg = total / zones
#             max_line = (avg + max(values)) / 2
#             min_line = (avg + min(values)) / 2
#         else:
#             total = avg = max_line = min_line = 0

#         # 4. Assign zone-based color scheme (Aapka code - Koi badlaav nahi)
#         color_map = {}
#         if divisions == 8:
#             color_map = {
#                 "N": "blue", "NE": "blue", "E": "green", "SE": "red",
#                 "S": "red", "SW": "yellow", "W": "grey", "NW": "grey"
#             }
#         elif divisions == 16:
#             for d in direction_counts:
#                 if d in ["NNW", "N", "NNE", "NE"]:
#                     color_map[d] = "blue"
#                 elif d in ["ENE", "E", "ESE"]:
#                     color_map[d] = "green"
#                 elif d in ["SE", "SSE", "S"]:
#                     color_map[d] = "red"
#                 elif d in ["SW", "SSW"]:
#                     color_map[d] = "yellow"
#                 elif d in ["WSW", "W", "WNW", "NW"]:
#                     color_map[d] = "grey"
#                 else:
#                     color_map[d] = "black"
#         elif divisions == 32:
#             for d in direction_counts:
#                 if d in ["N5", "N6", "N7", "N8", "E1", "N2", "N3", "N4"]:
#                     color_map[d] = "blue"
#                 elif d in ["E2", "E3", "E4", "E5", "E6", "E7"]:
#                     color_map[d] = "green"
#                 elif d in ["E8", "S1", "S2", "S3", "S4", "S5"]:
#                     color_map[d] = "red"
#                 elif d in ["S6", "S7", "S8", "W1"]:
#                     color_map[d] = "yellow"
#                 elif d in ["W2", "W3", "W4", "W5", "W6", "W7", "W8", "N1"]:
#                     color_map[d] = "grey"
#                 else:
#                     color_map[d] = "black"
#             pass
#         bar_colors = [color_map.get(k, "black") for k in direction_counts.keys()]

#         # 5. Plot Graph (Aapka code - Koi badlaav nahi)
#         fig, ax = plt.subplots(figsize=(10, 4))
#         ax.bar(direction_counts.keys(), direction_counts.values(), color=bar_colors)
#         ax.axhline(avg, color='red', linestyle='--', label='AVG AREA')
#         ax.axhline(max_line, color='purple', linestyle='--', label='MAX LINE')
#         ax.axhline(min_line, color='green', linestyle='--', label='MIN LINE')
#         ax.set_title("Zone-wise Area Distribution with Threshold Indicators")
#         ax.set_xlabel("Zone")
#         ax.set_ylabel("Area (sq units)")
#         ax.legend()
#         ax.grid(True, linestyle='--', alpha=0.6)
#         plt.xticks(rotation=45, ha="right")
#         plt.tight_layout()

#         graph_buffer = io.BytesIO()
#         plt.savefig(graph_buffer, format='PNG')
#         graph_buffer.seek(0)
#         plt.close(fig) # Graph banane ke baad figure ko close karna achhi practice hai

#         return HttpResponse(graph_buffer, content_type="image/png")

#     # Agar POST request na ho to yeh message dikhana
#     return HttpResponse("This endpoint requires a POST request.", status=405)


# download_blueprint old working code/..........................................................................................................................
# @csrf_exempt
# def download_blueprint(request):
#     """
#     Generates a zone-wise area bar chart image (PNG) using session grid_points and centroid info.

#     Behavior:
#     - Reads grid_points (flexible formats), compass_center, divisions, compass_rotation from session.
#     - Uses session['centroid_radius'] if present to only count points within that radius from compass_center.
#       Otherwise counts all grid_points.
#     - Maps points to sectors with 0° = North (clockwise), applies compass_rotation offset.
#     - Uses exact label orders you requested for 8 / 16 / 32.
#     - Returns PNG image of the bar chart.
#     """
#     try:
#         if request.method != "POST":
#             return HttpResponse("This endpoint requires a POST request.", status=405)

#         grid_points = request.session.get("grid_points", [])
#         compass_center = request.session.get("compass_center", [0, 0])
#         divisions = int(request.session.get("divisions", 8))
#         compass_rotation = float(request.session.get("compass_rotation", 0))
#         # Optional: radius up to which to count boxes (in same units as grid_points coords)
#         centroid_radius = request.session.get("centroid_radius", None)

#         # ensure compass_center numeric tuple
#         cx, cy = float(compass_center[0]), float(compass_center[1])

#         # --- Normalize grid_points into list of (x, y, inside_flag) ---
#         normalized = []
#         for p in grid_points:
#             # support different formats:
#             # tuple/list: (x,y) or (x,y,inside)
#             # dict: {"x":..,"y":.., "inside": True/False (optional)}
#             if isinstance(p, dict):
#                 x = float(p.get("x", p.get("0", 0)))
#                 y = float(p.get("y", p.get("1", 0)))
#                 inside = p.get("inside", True)
#             elif isinstance(p, (list, tuple)):
#                 if len(p) >= 3:
#                     x, y, inside = float(p[0]), float(p[1]), bool(p[2])
#                 else:
#                     x, y, inside = float(p[0]), float(p[1]), True
#             else:
#                 # can't parse: skip
#                 continue
#             normalized.append((x, y, inside))

#         # If no points available, still return an empty chart
#         if not normalized:
#             direction_counts = {}
#             labels = []
#         else:
#             # --- decide radius cutoff ---
#             if centroid_radius is not None:
#                 try:
#                     cutoff = float(centroid_radius)
#                 except Exception:
#                     cutoff = None
#             else:
#                 cutoff = None

#             # if cutoff not provided, include all points (or optionally compute max)
#             # We'll include all points if cutoff is None.

#             # --- label sequences as requested by user ---
#             labels_8 = ["N", "NE", "E", "SE", "S", "SW", "W", "NW"]
#             labels_16 = ["N","NNE","NE","ENE","E","ESE","SE","SSE","S","SSW","SW","WSW","W","WNW","NW","NNW"]
#             labels_32 = [
#                 "N5","N6","N7","N8",
#                 "E1","E2","E3","E4","E5","E6","E7","E8",
#                 "S1","S2","S3","S4","S5","S6","S7","S8",
#                 "W1","W2","W3","W4","W5","W6","W7","W8",
#                 "N1","N2","N3","N4"
#             ]

#             if divisions == 8:
#                 labels = labels_8
#             elif divisions == 16:
#                 labels = labels_16
#             elif divisions == 32:
#                 labels = labels_32
#             else:
#                 # fallback: generate generic labels
#                 sector_size = 360.0 / divisions
#                 labels = [f"S{i+1}" for i in range(divisions)]

#             # initialize counts
#             direction_counts = {lab: 0 for lab in labels}

#             # helper: compute angle where 0 = North, increasing clockwise
#             import math
#             def angle_from_center(x, y):
#                 dx = x - cx
#                 dy = y - cy
#                 # atan2 returns angle from x-axis (east) with positive counterclockwise.
#                 # We want 0 = North and clockwise positive.
#                 # One way: compute bearing = (90 - math.degrees(atan2(dy, dx))) % 360
#                 # But to ensure 0 = North and clockwise:
#                 bearing = (90.0 - math.degrees(math.atan2(dy, dx))) % 360.0
#                 return bearing

#             sector_size = 360.0 / float(divisions)

#             for (x, y, inside) in normalized:
#                 if not inside:
#                     continue
#                 # apply cutoff radius if given
#                 if cutoff is not None:
#                     dist = ((x - cx)**2 + (y - cy)**2) ** 0.5
#                     if dist > cutoff:
#                         continue

#                 ang = angle_from_center(x, y)
#                 # apply compass rotation offset (positive rotates clockwise)
#                 ang = (ang + float(compass_rotation)) % 360.0

#                 # center sectors around principal bearings: add half-sector to align
#                 # so that label at index 0 corresponds to angle near 0
#                 idx = int(((ang + sector_size / 2.0) % 360.0) // sector_size) % divisions

#                 # map idx -> label (for provided label lists)
#                 if idx < len(labels):
#                     lab = labels[idx]
#                 else:
#                     lab = labels[idx % len(labels)]
#                 direction_counts[lab] = direction_counts.get(lab, 0) + 1

#         # --- Prepare color mapping similar to your rules (keeps style) ---
#         bar_colors = []
#         # for 8 / 16 / 32 keep mapping you specified
#         if divisions == 8:
#             # Using colors closer to your first image but you can change to lighter shades
#             color_map = { "N": "lightblue", "NE": "lightblue", "E": "seagreen", "SE": "red", "S": "salmon", "SW": "yellow", "W": "lightgrey", "NW": "darkgrey" }
#             # If you want to strictly follow the first image's colors:
#             # color_map = { "N": "blue", "NE": "blue", "E": "green", "SE": "red", "S": "red", "SW": "yellow", "W": "grey", "NW": "grey" }

#         elif divisions == 16:
#             for d in direction_areas:
#                 if d in ["NNW", "N", "NNE", "NE"]: color_map[d] = "lightblue"
#                 elif d in ["ENE", "E", "ESE"]: color_map[d] = "seagreen"
#                 elif d in ["SE", "SSE", "S"]: color_map[d] = "salmon"
#                 elif d in ["SW", "SSW"]: color_map[d] = "yellow"
#                 else: color_map[d] = "lightgrey"
#         elif divisions == 32:
#             for d in direction_areas:
#                 if d in ["N5", "N6", "N7", "N8", "E1", "N2", "N3", "N4"]: color_map[d] = "lightblue"
#                 elif d in ["E2", "E3", "E4", "E5", "E6", "E7"]: color_map[d] = "seagreen"
#                 elif d in ["E8", "S1", "S2", "S3", "S4", "S5"]: color_map[d] = "salmon"
#                 elif d in ["S6", "S7", "S8", "W1"]: color_map[d] = "yellow"
#                 else: color_map[d] = "lightgrey"
#         else:
#              # Default color if divisions is not 8, 16, or 32
#             color_map = {key: "grey" for key in labels}

#         bar_colors = [color_map.get(k, "black") for k in direction_counts.keys()]

#         # --- compute some threshold lines for nicer plot (optional) ---
#         values = list(direction_counts.values())
#         if values:
#             total = sum(values)
#             zones = len(values)
#             avg = total / zones if zones else 0
#             max_line = (avg + max(values)) / 2.0
#             min_line = (avg + min(values)) / 2.0
#         else:
#             avg = max_line = min_line = 0

#         # --- plot using matplotlib ---
#         import matplotlib.pyplot as plt
#         fig, ax = plt.subplots(figsize=(10, 4))
#         keys = list(direction_counts.keys())
#         vals = [direction_counts[k] for k in keys]
#         ax.bar(keys, vals, color=bar_colors)
#         if any(values):
#             ax.axhline(avg, color='red', linestyle='--', label='AVG')
#             ax.axhline(max_line, color='purple', linestyle='--', label='MAX-LINE')
#             ax.axhline(min_line, color='green', linestyle='--', label='MIN-LINE')
#             ax.legend()
#         ax.set_title("Zone-wise Box Count Distribution")
#         ax.set_xlabel("Zone")
#         ax.set_ylabel("Count")
#         ax.grid(True, linestyle='--', alpha=0.5)
#         plt.xticks(rotation=45, ha="right")
#         plt.tight_layout()

#         graph_buffer = io.BytesIO()
#         plt.savefig(graph_buffer, format='PNG', dpi=150)
#         graph_buffer.seek(0)
#         plt.close(fig)

#         return HttpResponse(graph_buffer.getvalue(), content_type="image/png")
#     except Exception:
#         # log on server console (helpful)
#         print("ERROR generating graph:")
#         print(traceback.format_exc())
#         return HttpResponse(status=500)


# old working code----------------------------------------------------------------------------------------------------------------------------------------
# @csrf_exempt
# def generate_graph_data_view(request, project_id):
#     if request.method == 'POST':
#         try:
#             data = json.loads(request.body)
#             divisions = data.get('divisions')

#             if not divisions:
#                 return JsonResponse({'error': 'Divisions not provided'}, status=400)

#             compass_rotation = request.session.get("compass_rotation", 0)
#             grid_points = request.session.get("grid_points", [])
#             compass_center = request.session.get("compass_center", [])

#             if not grid_points or not compass_center:
#                 return JsonResponse({'error': 'Grid or compass data not found in session. Please recalculate.'}, status=400)

#             image_bytes_io = generate_graph_image(grid_points, compass_center, divisions, compass_rotation)


#             if image_bytes_io:
#                 image_base64 = base64.b64encode(image_bytes_io.read()).decode('utf-8')

#                 return JsonResponse({
#                     'graph_image_base64': image_base64
#                 })
#             else:
#                 return JsonResponse({'error': 'The generate_graph_image function failed and returned None.'}, status=500)

#         except Exception as e:
#             print(f"ERROR in generate_graph_data_view: {e}") 
#             return JsonResponse({'error': f'An unexpected error occurred on the server: {str(e)}'}, status=500)

#     return JsonResponse({'error': 'Invalid request method'}, status=405)

# def generate_graph_image(grid_points, compass_center, divisions, compass_rotation):
#     try:
#         direction_counts = calculate_directional_areas(
#             [tuple(p) for p in grid_points], tuple(compass_center), divisions, compass_rotation
#         )

#         values = list(direction_counts.values())
#         if divisions > 0 and values:
#             total = sum(values)
#             avg = total / divisions
#             max_line = (avg + max(values)) / 2
#             min_line = (avg + min(values)) / 2
#         else:
#             avg, max_line, min_line = 0, 0, 0

#         color_map = {}

#         if divisions == 8:
#             color_map = { "N": "blue", "NE": "blue", "E": "green", "SE": "red", "S": "red", "SW": "yellow", "W": "grey", "NW": "grey" }
#         elif divisions == 16:
#             for d in direction_counts:
#                 if d in ["NNW", "N", "NNE", "NE"]: color_map[d] = "blue"
#                 elif d in ["ENE", "E", "ESE"]: color_map[d] = "green"
#                 elif d in ["SE", "SSE", "S"]: color_map[d] = "red"
#                 elif d in ["SW", "SSW"]: color_map[d] = "yellow"
#                 else: color_map[d] = "grey"
#         elif divisions == 32:
#             for d in direction_counts:
#                 if d in ["N5", "N6", "N7", "N8", "E1", "N2", "N3", "N4"]: color_map[d] = "blue"
#                 elif d in ["E2", "E3", "E4", "E5", "E6", "E7"]: color_map[d] = "green"
#                 elif d in ["E8", "S1", "S2", "S3", "S4", "S5"]: color_map[d] = "red"
#                 elif d in ["S6", "S7", "S8", "W1"]: color_map[d] = "yellow"
#                 else: color_map[d] = "grey"

#         bar_colors = [color_map.get(k, "black") for k in direction_counts.keys()]


#         fig, ax = plt.subplots(figsize=(10, 4))
#         ax.bar(direction_counts.keys(), direction_counts.values(), color=bar_colors)
#         ax.axhline(avg, color='red', linestyle='--', label='AVG AREA')
#         ax.axhline(max_line, color='purple', linestyle='--', label='MAX LINE')
#         ax.axhline(min_line, color='green', linestyle='--', label='MIN LINE')
#         ax.set_title(f"Zone-wise Area Distribution ({divisions} Parts)")
#         ax.set_xlabel("Zone")
#         ax.set_ylabel("Area (sq units)")
#         ax.legend()
#         ax.grid(True, linestyle='--', alpha=0.6)
#         plt.xticks(rotation=45, ha="right")
#         plt.tight_layout()

#         graph_buffer = io.BytesIO()
#         plt.savefig(graph_buffer, format='PNG')
#         graph_buffer.seek(0)
#         plt.close(fig)
#         return graph_buffer
#     except Exception as e:
#         print(f"Error generating graph for {divisions} divisions: {e}")
#         return None
# old working code ends here--------------------------------------------------------------------------------------------------------------------------------


# def download_word_blueprint(request, project_id):
#     try:
#         data = {}
#         compass_degree = 'N/A'
#         degree_north = 'N/A'
#         degree_east = 'N/A'
#         degree_south = 'N/A'
#         degree_west = 'N/A'
#         if request.body:
#             try:
#                 data = json.loads(request.body)
#                 value = data.get('compass_degree')
#                 degree_north = data.get('degree_north', 'N/A')
#                 degree_east = data.get('degree_east', 'N/A')
#                 degree_south = data.get('degree_south', 'N/A')
#                 degree_west = data.get('degree_west', 'N/A')
#                 if value is not None and str(value).strip():
#                     compass_degree = value
#             except json.JSONDecodeError:
#                 print("No valid JSON in request body. Using defaults.")

#         project = get_object_or_404(Project, pk=project_id)

#         user_name = request.user.username
#         user_email = request.user.email

#         document = Document()

#         document.add_heading('User and Project Details', level=1)

#         user_table = document.add_table(rows=2, cols=2)
#         user_table.style = 'Table Grid'
#         user_table.cell(0, 0).text = 'Username'
#         user_table.cell(0, 1).text = user_name
#         user_table.cell(1, 0).text = 'Email'
#         user_table.cell(1, 1).text = user_email

#         for row in user_table.rows:
#             row.cells[0].paragraphs[0].runs[0].font.bold = True

#         document.add_paragraph() 

#         project_table = document.add_table(rows=4, cols=2)
#         project_table.style = 'Table Grid'
#         project_table.cell(0, 0).text = 'Project Name'
#         project_table.cell(0, 1).text = project.name
#         project_table.cell(1, 0).text = 'Description'
#         project_table.cell(1, 1).text = project.description
#         project_table.cell(2, 0).text = 'Category'
#         project_table.cell(2, 1).text = project.category
#         project_table.cell(3, 0).text = 'Status'
#         project_table.cell(3, 1).text = project.status

#         for row in project_table.rows:
#             row.cells[0].paragraphs[0].runs[0].font.bold = True

#         document.add_page_break()

#         document.add_heading('Blueprint Analysis Report', level=1)   

#         # Fetch session data
#         compass_rotation = request.session.get("compass_rotation", 0)
#         grid_points = request.session.get("grid_points", [])
#         compass_center = request.session.get("compass_center", [0, 0])

#         def add_image_to_doc(doc, heading, image_field, add_page_break_after=True):
#             doc.add_heading(heading, level=2)
#             if image_field and hasattr(image_field, 'path'):
#                 try:
#                     doc.add_picture(image_field.path, width=Inches(6.0))
#                 except Exception:
#                     doc.add_paragraph(f"Image file not found or is invalid.")
#             else:
#                 doc.add_paragraph("Image not available.")

#             if add_page_break_after:
#                 doc.add_page_break()

#         add_image_to_doc(document, 'Original Blueprint', project.blueprint)
#         add_image_to_doc(document, 'Centroid Calculated Image', project.centroid_image)

#         add_image_to_doc(document, 'Compass Set Image', project.compass_image, add_page_break_after=False)
#         document.add_paragraph()
#         compass_table = document.add_table(rows=1, cols=2)
#         compass_table.style = 'Table Grid'

#         hdr_cells = compass_table.rows[0].cells
#         hdr_cells[0].text = 'Direction'
#         hdr_cells[1].text = 'Stopped At Degree'

#         for cell in hdr_cells:
#             cell.paragraphs[0].runs[0].font.bold = True

#         row_cells = compass_table.add_row().cells
#         row_cells[0].text = "North"
#         row_cells[1].text = f"{degree_north}°"
#         document.add_page_break()

#         # --- 8 Parts ---
#         add_image_to_doc(document, 'Image Divided into 8 Parts', project.divided_8_image, add_page_break_after=False)
#         graph8 = generate_graph_image(grid_points, compass_center, 8, compass_rotation)
#         if graph8:
#             document.add_paragraph("Graph for 8 Parts:")
#             document.add_picture(graph8, width=Inches(6.0))
#         document.add_page_break()

#         # --- 16 Parts ---
#         add_image_to_doc(document, 'Image Divided into 16 Parts', project.divided_16_image, add_page_break_after=False)
#         graph16 = generate_graph_image(grid_points, compass_center, 16, compass_rotation)
#         if graph16:
#             document.add_paragraph("Graph for 16 Parts:")
#             document.add_picture(graph16, width=Inches(6.0))
#         document.add_page_break()

#         # --- 32 Parts (અહીં છેલ્લો પેજ બ્રેક નહીં આવે) ---
#         add_image_to_doc(document, 'Image Divided into 32 Parts', project.divided_32_image, add_page_break_after=False)
#         graph32 = generate_graph_image(grid_points, compass_center, 32, compass_rotation)
#         if graph32:
#             document.add_paragraph("Graph for 32 Parts:")
#             document.add_picture(graph32, width=Inches(6.0))

#         # Save Word file in memory
#         doc_io = io.BytesIO()
#         document.save(doc_io)
#         doc_io.seek(0)

#         response = HttpResponse(
#             doc_io.read(),
#             content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#         )
#         response['Content-Disposition'] = f'attachment; filename="analysis_{project.id}.docx"'
#         return response

#     except Exception as e:
#         print("ERROR IN WORD GENERATION:")
#         print(traceback.format_exc())
#         return HttpResponse(f"<h1>An Error Occurred</h1><p>{e}</p><pre>{traceback.format_exc()}</pre>", status=500)

# new function-----------------------------------------------------------------------------------
def download_word_blueprint(request, project_id):
    try:
        # 1. READ PAYLOAD (Degrees & Images from Frontend)
        data = {}
        compass_degree = 'N/A'
        degree_north = 'N/A'
        degree_east = 'N/A'
        degree_south = 'N/A'
        degree_west = 'N/A'

        data = json.loads(request.body)
        marked_rooms = data.get('marked_rooms', [])
        compass_degree = data.get('compass_degree', 0)

        if request.body:
            try:
                data = json.loads(request.body)
                value = data.get('compass_degree')
                degree_north = data.get('degree_north', 'N/A')
                degree_east = data.get('degree_east', 'N/A')
                degree_south = data.get('degree_south', 'N/A')
                degree_west = data.get('degree_west', 'N/A')
                if value is not None and str(value).strip():
                    compass_degree = value
            except json.JSONDecodeError:
                pass

        project = get_object_or_404(Project, pk=project_id)
        user_name = request.user.username

        # 2. DOCUMENT SETUP (Styles & Reduced Margins)
        document = Document()


        # --- Styles ---
        styles = document.styles
        if 'CustomHeading1' not in [s.name for s in styles]:
            h = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(12)
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.font.name = 'Arial Black'
            h.font.size = Pt(28)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        if 'CustomNormal' not in [s.name for s in styles]:
            n = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
            n.paragraph_format.space_before = Pt(6)
            n.paragraph_format.space_after = Pt(6)
            n.font.name = 'Calibri'
            n.font.size = Pt(12)
            n.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # --- Page Setup (Maximize Printable Area) ---
        section = document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

        # --- Blue Page Border ---
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f"w:{border_name}")
            border_el.set(qn('w:val'), 'single')
            border_el.set(qn('w:sz'), '12')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), '4F81BD')
            pgBorders.append(border_el)
        sectPr.append(pgBorders)

        # --- Pagination Footer ---
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Page ")
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.add_text(" of ")
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = "NUMPAGES"
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)

        # 3. COVER PAGE CONTENT
        document.add_paragraph("Blueprint Analysis Report", style='CustomHeading1')

        title = document.add_paragraph(project.name)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if title.runs:
            title.runs[0].font.size = Pt(34)
            title.runs[0].font.bold = True
            title.runs[0].font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

        document.add_paragraph(f"Prepared for: {user_name}", style='CustomNormal')
        document.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", style='CustomNormal')

        document.add_paragraph()
        subhead = document.add_paragraph("Project Details")
        subhead.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if subhead.runs:
            subhead.runs[0].font.bold = True
            subhead.runs[0].font.size = Pt(14)

        document.add_paragraph(f"Project Name: {project.name}", style='CustomNormal')
        document.add_paragraph(f"Description: {project.description}", style='CustomNormal')
        document.add_paragraph(f"Category: {project.category}", style='CustomNormal')
        document.add_paragraph(f"Status: {project.status}", style='CustomNormal')

        document.add_page_break()

        # 4. SMART IMAGE HELPER (FIXED: DATABASE PRIORITY)
        def add_smart_image(doc, heading, db_image_field=None, json_key_name=None, add_page_break_after=True):
            doc.add_heading(heading, level=2)

            image_added = False

            # PRIORITY 1: Database Image
            if db_image_field and hasattr(db_image_field, 'path'):
                try:
                    if os.path.exists(db_image_field.path):
                        doc.add_picture(db_image_field.path, width=Inches(7.5))
                        image_added = True
                except Exception as e:
                    pass

            # PRIORITY 2: Frontend Data (Fallback only)
            if not image_added and json_key_name and data.get(json_key_name):
                try:
                    b64_data = data.get(json_key_name)
                    if "," in b64_data:
                        b64_data = b64_data.split(',')[-1]
                    image_stream = io.BytesIO(base64.b64decode(b64_data))
                    doc.add_picture(image_stream, width=Inches(7.5))
                    image_added = True
                except Exception as e:
                    pass

            if not image_added:
                doc.add_paragraph("Image not available.")

            if add_page_break_after:
                doc.add_page_break()

        # 5. ADD IMAGES (With Database Priority)

        # Original & Centroid
        add_smart_image(document, 'Original Blueprint', db_image_field=project.blueprint)
        add_smart_image(document, 'Centroid Calculated Image', db_image_field=project.centroid_image)

        # Compass (DB Preferred)
        add_smart_image(document, 'Compass Set Image', db_image_field=project.compass_image,
                        json_key_name='compass_image', add_page_break_after=False)
        document.add_paragraph()

        # Data Table
        compass_table = document.add_table(rows=1, cols=2)
        compass_table.style = "Light Shading Accent 1"
        hdr = compass_table.rows[0].cells
        hdr[0].text = 'Direction'
        hdr[1].text = 'Stopped At Degree'
        for c in hdr:
            for r in c.paragraphs[0].runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row = compass_table.add_row().cells
        row[0].text = "North"
        row[1].text = f"{degree_north}°"
        for r in row[0].paragraphs[0].runs:
            r.font.bold = True
        document.add_page_break()

        # ---------------- GRIDS (8, 16, 32) ----------------

        # >>> 8 Parts
        add_smart_image(document, 'Image Divided into 8 Parts', db_image_field=project.divided_8_image,
                        json_key_name='divided_8_image', add_page_break_after=False)
        document.add_page_break()

        # Graph 8
        if data.get('graph_8_image'):
            doc_heading = document.add_heading("Graph for 8 Parts", level=2)
            try:
                b64 = data.get('graph_8_image').split(',')[-1]
                document.add_picture(io.BytesIO(base64.b64decode(b64)), width=Inches(7.5))
            except:
                document.add_paragraph("Graph image error.")
            document.add_page_break()
        else:
            grid_points = request.session.get("grid_points", [])
            compass_center = request.session.get("compass_center", [0, 0])
            if grid_points:
                g8 = generate_graph_image(grid_points, compass_center, 8, request.session.get("compass_rotation", 0))
                if g8:
                    document.add_heading("Graph for 8 Parts", level=2)
                    document.add_picture(g8, width=Inches(7.5))
                    document.add_page_break()
                else:
                    document.add_page_break()
            else:
                document.add_page_break()

        # >>> 16 Parts
        add_smart_image(document, 'Image Divided into 16 Parts', db_image_field=project.divided_16_image,
                        json_key_name='divided_16_image', add_page_break_after=False)
        document.add_page_break()

        # Graph 16
        if data.get('graph_16_image'):
            document.add_heading("Graph for 16 Parts", level=2)
            try:
                b64 = data.get('graph_16_image').split(',')[-1]
                document.add_picture(io.BytesIO(base64.b64decode(b64)), width=Inches(7.5))
            except:
                pass
            document.add_page_break()
        else:
            grid_points = request.session.get("grid_points", [])
            compass_center = request.session.get("compass_center", [0, 0])
            if grid_points:
                g16 = generate_graph_image(grid_points, compass_center, 16, request.session.get("compass_rotation", 0))
                if g16:
                    document.add_heading("Graph for 16 Parts", level=2)
                    document.add_picture(g16, width=Inches(7.5))
                    document.add_page_break()
                else:
                    document.add_page_break()
            else:
                document.add_page_break()

        # >>> 32 Parts
        add_smart_image(document, 'Image Divided into 32 Parts', db_image_field=project.divided_32_image,
                        json_key_name='divided_32_image', add_page_break_after=False)
        document.add_page_break()

        # Graph 32
        if data.get('graph_32_image'):
            document.add_heading("Graph for 32 Parts", level=2)
            try:
                b64 = data.get('graph_32_image').split(',')[-1]
                document.add_picture(io.BytesIO(base64.b64decode(b64)), width=Inches(7.5))
            except:
                pass
        else:
            grid_points = request.session.get("grid_points", [])
            compass_center = request.session.get("compass_center", [0, 0])
            if grid_points:
                g32 = generate_graph_image(grid_points, compass_center, 32, request.session.get("compass_rotation", 0))
                if g32:
                    document.add_heading("Graph for 32 Parts", level=2)
                    document.add_picture(g32, width=Inches(7.5))

        # ── NEW: Room Vastu Analysis page (bilkul end mein) ──
        if marked_rooms:
            room_analysis = get_ai_room_analysis(marked_rooms, compass_degree)
            add_room_analysis_page(document, marked_rooms, room_analysis)

        # 6. RETURN FILE
        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0)

        response = HttpResponse(
            doc_io.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="analysis_{project.id}.docx"'
        return response

    except Exception as e:
        print("ERROR IN WORD GENERATION:")
        print(traceback.format_exc())
        if request.headers.get('x-requested-with') == 'XMLHttpRequest' or request.content_type == 'application/json':
            return JsonResponse({'error': str(e)}, status=500)
        return HttpResponse(f"<h1>An Error Occurred</h1><p>{e}</p>", status=500)


# ═══════════════════════════════════════════════════════════
#   HELPER 1 — AI se room-wise Vastu analysis lao
# ═══════════════════════════════════════════════════════════
def get_ai_room_analysis(marked_rooms, compass_degree):
    from openai import OpenAI
    from django.conf import settings

    rooms_text = "\n".join([
        f"- {r['name']} in {r['direction']} direction"
        for r in marked_rooms
    ])

    prompt = f"""
You are an expert Vastu Shastra consultant.

Compass orientation: {compass_degree}°

Rooms marked in this blueprint:
{rooms_text}

For each room, provide Vastu analysis. Return ONLY valid JSON, no extra text:

{{
  "rooms": [
    {{
      "name": "<room name>",
      "direction": "<direction>",
      "vastu_status": "Excellent / Good / Average / Poor",
      "vastu_score": <0-100>,
      "why": "<why this placement is good or bad according to Vastu>",
      "effects": "<what effects this placement has on residents>",
      "remedies": ["<remedy 1>", "<remedy 2>"]
    }}
  ],
  "overall_room_score": <0-100>,
  "overall_room_summary": "<overall assessment of room placements>"
}}
"""

    try:
        client = OpenAI(
            api_key=settings.GROQ_API_KEY,
            base_url="https://api.groq.com/openai/v1"
        )
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "Return ONLY valid JSON."},
                {"role": "user",   "content": prompt}
            ],
            temperature=0.6
        )
        raw = response.choices[0].message.content.strip()
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            import re
            match = re.search(r'\{.*\}', raw, re.DOTALL)
            return json.loads(match.group()) if match else {"rooms": [], "overall_room_summary": ""}
    except Exception as e:
        print(f"AI Room Analysis Error: {e}")
        return {"rooms": [], "overall_room_score": 0, "overall_room_summary": "AI analysis unavailable."}


# ═══════════════════════════════════════════════════════════
#   HELPER 2 — Word doc mein Room Analysis page add karo
# ═══════════════════════════════════════════════════════════
def add_room_analysis_page(document, marked_rooms, room_analysis):
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # ── Page break → new page ──
    document.add_page_break()

    # ── Page title ──
    title = document.add_heading('Room-wise Vastu Analysis', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    document.add_paragraph()

    # ── Overall score & summary ──
    overall_score   = room_analysis.get('overall_room_score', 0)
    overall_summary = room_analysis.get('overall_room_summary', '')

    score_para = document.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_run = score_para.add_run(f"Overall Room Placement Score:  {overall_score} / 100")
    score_run.bold      = True
    score_run.font.size = Pt(14)
    if overall_score >= 75:
        score_run.font.color.rgb = RGBColor(0x15, 0x80, 0x3d)
    elif overall_score >= 50:
        score_run.font.color.rgb = RGBColor(0x92, 0x40, 0x0e)
    else:
        score_run.font.color.rgb = RGBColor(0x99, 0x1b, 0x1b)

    if overall_summary:
        summary_para = document.add_paragraph(overall_summary)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if summary_para.runs:
            summary_para.runs[0].font.size = Pt(11)

    document.add_paragraph()

    # ── Status color mapping ──
    status_colors = {
        'Excellent': RGBColor(0x15, 0x80, 0x3d),
        'Good':      RGBColor(0x16, 0x65, 0x34),
        'Average':   RGBColor(0x92, 0x40, 0x0e),
        'Poor':      RGBColor(0x99, 0x1b, 0x1b),
    }

    ai_rooms = room_analysis.get('rooms', [])

    # ── Per-room section ──
    for room_data in ai_rooms:
        name      = room_data.get('name', '')
        direction = room_data.get('direction', '')
        status    = room_data.get('vastu_status', 'Average')
        score     = room_data.get('vastu_score', 0)
        why       = room_data.get('why', '')
        effects   = room_data.get('effects', '')
        remedies  = room_data.get('remedies', [])

        # Room heading
        room_heading = document.add_heading(f"{name}  ({direction})", level=2)
        if room_heading.runs:
            room_heading.runs[0].font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

        # Status + score line
        status_para = document.add_paragraph()
        status_run  = status_para.add_run(f"Vastu Status: {status}   |   Score: {score} / 100")
        status_run.bold      = True
        status_run.font.size = Pt(11)
        status_run.font.color.rgb = status_colors.get(status, RGBColor(0x37, 0x41, 0x51))

        # Why
        if why:
            why_para = document.add_paragraph()
            why_label = why_para.add_run("Vastu Analysis:  ")
            why_label.bold = True
            why_label.font.size = Pt(10)
            why_text = why_para.add_run(why)
            why_text.font.size = Pt(10)

        # Effects
        if effects:
            eff_para = document.add_paragraph()
            eff_label = eff_para.add_run("Effects on Residents:  ")
            eff_label.bold = True
            eff_label.font.size = Pt(10)
            eff_text = eff_para.add_run(effects)
            eff_text.font.size = Pt(10)

        # Remedies
        if remedies:
            rem_heading_para = document.add_paragraph()
            rem_label = rem_heading_para.add_run("Remedies:")
            rem_label.bold      = True
            rem_label.font.size = Pt(10)
            for remedy in remedies:
                bullet = document.add_paragraph(style='List Bullet')
                bullet_run = bullet.add_run(remedy)
                bullet_run.font.size = Pt(10)

        # Separator
        sep = document.add_paragraph("─" * 80)
        if sep.runs:
            sep.runs[0].font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            sep.runs[0].font.size = Pt(8)

        document.add_paragraph()

    # ── Footer note ──
    document.add_paragraph()
    note_para = document.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note_para.add_run("Generated by AI Vastu Analysis System  |  Room placement analysis based on Vastu Shastra principles")
    note_run.font.size = Pt(8)
    note_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    note_run.italic = True


@login_required
def ajax_load_projects(request):
    # get projects of logged in users only
    projects = Project.objects.filter(user=request.user)

    # 1. search by name
    search_term = request.GET.get('search', '').strip()
    if search_term:
        projects = projects.filter(name__icontains=search_term)

    # 2. filter by status & category
    status = request.GET.get('status', '')
    if status:
        projects = projects.filter(status=status)

    category = request.GET.get('category', '')
    if category:
        projects = projects.filter(category=category)

    # 3. sorting
    sort_by = request.GET.get('sort', 'date_new')
    if sort_by == 'date_new':
        projects = projects.order_by('-created_at')  # 'created_at' model field
    elif sort_by == 'date_old':
        projects = projects.order_by('created_at')
    elif sort_by == 'name_asc':
        projects = projects.order_by('name')
    elif sort_by == 'name_desc':
        projects = projects.order_by('-name')
    project_list = list(projects.values('id', 'name', 'description', 'status', 'category'))

    return JsonResponse({'projects': project_list})


@login_required
def ajax_delete_projects(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            project_ids = data.get('project_ids', [])

            if not project_ids:
                return JsonResponse({'status': 'error', 'message': 'No project IDs provided'}, status=400)

            projects_to_delete = Project.objects.filter(user=request.user, id__in=project_ids)
            count = projects_to_delete.count()
            projects_to_delete.delete()

            return JsonResponse({'status': 'success', 'message': f'{count} projects deleted successfully.'})
        except json.JSONDecodeError:
            return JsonResponse({'status': 'error', 'message': 'Invalid JSON'}, status=400)

    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=405)


def logout_view(request):
    logout(request)
    request.session.flush()
    return redirect('login')


# @csrf_exempt
# def analyze_grid(request):
#     if request.method == "POST":
#         data = json.loads(request.body)
#
#         grid_points = data.get("grid_data", [])
#         center = data.get("compass_center", [0, 0])
#         divisions = int(data.get("divisions", 8))
#         compass_rotation = float(data.get("compass_rotation", 0))
#
#         # ✅ Save ALL required data to session
#         request.session["grid_points"] = grid_points
#         request.session["compass_center"] = center
#         request.session["divisions"] = divisions
#         request.session["compass_rotation"] = compass_rotation
#
#         # ✅ Also save centroid data if available
#         if 'centroid' in data:
#             request.session["centroid"] = data['centroid']
#
#         print(f"DEBUG: Saved to session - divisions: {divisions}, rotation: {compass_rotation}")
#
#         direction_counts = calculate_directional_areas(
#             [tuple(p) for p in grid_points],
#             tuple(center),
#             divisions,
#             compass_rotation
#         )
#
#         print("\n🔍 Direction-wise Block Counts:")
#         for k, v in direction_counts.items():
#             print(f"{k}: {v}")
#
#         return JsonResponse(direction_counts)


# admin page
# ✅ Static admin credentials
ADMIN_USERNAME = 'admin'
ADMIN_PASSWORD = 'admin11'


def admin_login(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')

        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            request.session['admin_logged_in'] = True
            return redirect('admin_dashboard')
        else:
            return render(request, 'admin_login.html', {'error': 'Invalid username or password'})

    return render(request, 'admin_login.html')


# def admin_dashboard(request):
#     if not request.session.get('admin_logged_in'):
#         return redirect('admin_login')
#
#     # Subquery to get project_limit from UserProfile for each user
#     project_limit_subquery = UserProfile.objects.filter(user=OuterRef('pk')).values('project_limit')[:1]
#
#     # Annotate each user with their project_count and project_limit
#     users = User.objects.annotate(
#         project_count=Count('project'),
#         project_limit=Subquery(project_limit_subquery)
#     )
#
#     # projects = Project.object.select_related('user')
#     total_users = users.count()
#
#     return render(request, 'admin_dashboard.html', {
#         'users': users,
#         'total_users': total_users
#     })

def admin_dashboard(request):

    project_limit_subquery = UserProfile.objects.filter(user=OuterRef('pk')).values('project_limit')[:1]

    users = User.objects.annotate(
        project_count=Count('project'),
        project_limit=Subquery(project_limit_subquery)
    )

    total_users = users.count()

    # All projects with user info
    all_projects = Project.objects.select_related('user').order_by('-created_at')

    # All vastu reports with user + project info
    all_reports = VastuReport.objects.select_related('project', 'project__user').order_by('-created_at')
    total_projects = all_projects.count()
    # Payment data — user wise
    # Payment data
    all_paid_orders = PaymentOrder.objects.filter(
        status='paid'
    ).select_related('user').order_by('-created_at')

    from django.db.models import Sum
    total_revenue = PaymentOrder.objects.filter(status='paid').aggregate(
        total=Sum('amount')
    )['total'] or 0

    total_paying_users = PaymentOrder.objects.filter(
        status='paid'
    ).values('user').distinct().count()

    return render(request, 'admin_dashboard.html', {
        'users': users,
        'total_users': total_users,
        'all_projects': all_projects,
        'total_projects': total_projects,
        'all_paid_orders': all_paid_orders,
        'total_revenue': total_revenue,
        'total_paying_users': total_paying_users,
    })


def admin_user_detail(request, user_id):

    user = get_object_or_404(User, id=user_id)
    projects = Project.objects.filter(user=user).order_by('-created_at')
    reports = VastuReport.objects.filter(project__user=user).select_related('project', 'project__user').order_by(
        '-created_at')
    return render(request, 'admin_user_detail.html', {
        'selected_user': user,
        'projects': projects,
        'reports': reports,
    })

# def admin_user_detail(request, user_id):
#     if not request.session.get('admin_logged_in'):
#         return redirect('admin_login')
#     user = get_object_or_404(User, id=user_id)
#     projects = Project.objects.filter(user=user).order_by('-created_at')
#     reports = VastuReport.objects.filter(project__user=user).select_related('project').order_by('-created_at')
#     return render(request, 'admin_user_detail.html', {
#         'selected_user': user,
#         'projects': projects,
#         'reports': reports,
#     })

@csrf_exempt
def delete_project(request):
    if request.method == 'POST':
        project_id = request.POST.get('project_id')
        try:
            project = Project.objects.get(id=project_id)
            project.delete()
            messages.success(request, f'Project deleted successfully.')
        except Project.DoesNotExist:
            pass
    return redirect('admin_dashboard')

def admin_logout(request):
    request.session.flush()
    return redirect('admin_login')


@csrf_exempt
def delete_user(request):
    if request.method == 'POST':
        user_id = request.POST.get('user_id')
        try:
            user = User.objects.get(id=user_id)
            user.delete()
        except Users.DoesNotExist:
            pass
    return redirect('admin_dashboard')

